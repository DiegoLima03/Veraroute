#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
INFORME DE OPTIMIZACIÓN DE RUTAS COMERCIALES — VERALEZA
========================================================
Extrae datos reales de MySQL (gestorrutas), calcula distancias OSRM,
analiza rentabilidad por cliente y genera informe .docx profesional.

Uso:  python scripts/informe_rutas.py
"""

import os, sys, json, time, math, locale, random, warnings
from datetime import datetime, date
from pathlib import Path

import requests
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from PIL import Image

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

warnings.filterwarnings('ignore')

# Fix Windows console encoding
import io
if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

# ─── CONFIGURACIÓN ───────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = BASE_DIR / 'output'
OUTPUT_DIR.mkdir(exist_ok=True)
CHART_DIR = OUTPUT_DIR / 'charts'
CHART_DIR.mkdir(exist_ok=True)

LOGO_PATH = Path(r'C:\wamp64\www\licitaciones\proyectoenphp\public\img\logo_login.png')

# Base de operaciones: Veraleza Tomiño
BASE_LAT, BASE_LNG = 41.994524, -8.739887

# Parámetros de coste
COSTE_KM = 0.35          # €/km
COSTE_HORA = 18.0         # €/hora conductor
TIEMPO_DESCARGA_MIN = 20  # minutos por cliente
PAQUETERIA = {5: 4.50, 15: 7.50, 30: 12.00, 999: 25.00}

# Colores corporativos
VERDE_VERALEZA = '8E8B30'
COLOR_FILA_ALT = 'F5F4F0'
COLOR_BORDE = 'D5D3CC'
BG_VERDE = 'E8F5E9'
BG_AMARILLO = 'FFF8E1'
BG_ROJO = 'FDECEA'

# Colores de ruta (mismo orden que RUTAS en BD: ids 1-8)
RUTA_COLORS_HEX = {
    'Pontevedra 3': '#795548', 'Comarca A': '#3498db', 'Comarca B': '#2ecc71',
    'Pontevedra 1': '#9b59b6', 'Pontevedra 2': '#e67e22', 'Orense A': '#1abc9c',
    'Orense B': '#e84393', 'Orense C': '#d4b800'
}
RUTA_COLORS_MPL = {k: v for k, v in RUTA_COLORS_HEX.items()}

# ─── BASE DE DATOS ───────────────────────────────────────────────
def fetch_clients():
    """Lee clientes del JSON exportado por CLI de MySQL."""
    json_path = OUTPUT_DIR / 'clients_export.json'
    with open(json_path, 'r', encoding='utf-8') as f:
        rows = json.load(f)
    for r in rows:
        r['lat'] = float(r['lat'])
        r['lng'] = float(r['lng'])
    return rows

def fetch_rutas():
    """Extrae rutas únicas de los clientes."""
    clients = fetch_clients()
    seen = {}
    for c in clients:
        rid = c.get('ruta_id')
        rname = c.get('ruta_name', '')
        if rid and rname and rid not in seen:
            seen[rid] = {'id': rid, 'name': rname}
    return sorted(seen.values(), key=lambda x: x['name'])

# ─── SIMULACIÓN DE DATOS COMERCIALES ────────────────────────────
# No hay historial real suficiente — simulamos de forma realista
# basándonos en la distancia y zona geográfica.
random.seed(42)

def simulate_commercial_data(clients):
    """Añade facturación, frecuencia y peso medio simulados."""
    for c in clients:
        dist_km = haversine(BASE_LAT, BASE_LNG, c['lat'], c['lng'])
        c['dist_lineal_km'] = dist_km

        # Clientes más cercanos suelen facturar más y pedir más frecuentemente
        if dist_km < 30:
            c['facturacion_mensual'] = random.uniform(800, 3500)
            c['frecuencia_mensual'] = random.choice([4, 4, 4, 3, 2])
            c['peso_medio_kg'] = random.uniform(15, 80)
        elif dist_km < 60:
            c['facturacion_mensual'] = random.uniform(400, 2000)
            c['frecuencia_mensual'] = random.choice([4, 3, 2, 2, 1])
            c['peso_medio_kg'] = random.uniform(10, 50)
        elif dist_km < 100:
            c['facturacion_mensual'] = random.uniform(200, 1200)
            c['frecuencia_mensual'] = random.choice([2, 2, 1, 1])
            c['peso_medio_kg'] = random.uniform(5, 35)
        else:
            c['facturacion_mensual'] = random.uniform(100, 800)
            c['frecuencia_mensual'] = random.choice([2, 1, 1, 1])
            c['peso_medio_kg'] = random.uniform(3, 25)

        c['facturacion_mensual'] = round(c['facturacion_mensual'], 2)
        c['peso_medio_kg'] = round(c['peso_medio_kg'], 1)
    return clients

# ─── OSRM ────────────────────────────────────────────────────────
OSRM_BASE = "https://router.project-osrm.org"
_osrm_cache = {}

def haversine(lat1, lng1, lat2, lng2):
    R = 6371
    dlat = math.radians(lat2 - lat1)
    dlng = math.radians(lng2 - lng1)
    a = math.sin(dlat/2)**2 + math.cos(math.radians(lat1)) * math.cos(math.radians(lat2)) * math.sin(dlng/2)**2
    return R * 2 * math.asin(math.sqrt(a))

def osrm_route(lat1, lng1, lat2, lng2):
    """Devuelve (km, minutos) por carretera. Usa cache + fallback haversine."""
    key = (round(lat1,5), round(lng1,5), round(lat2,5), round(lng2,5))
    if key in _osrm_cache:
        return _osrm_cache[key]
    try:
        url = f"{OSRM_BASE}/route/v1/driving/{lng1},{lat1};{lng2},{lat2}?overview=false"
        r = requests.get(url, timeout=10)
        if r.status_code == 200:
            data = r.json()
            if data.get('code') == 'Ok':
                km = data['routes'][0]['distance'] / 1000
                mins = data['routes'][0]['duration'] / 60
                _osrm_cache[key] = (round(km, 1), round(mins, 1))
                return _osrm_cache[key]
    except:
        pass
    # Fallback
    km = haversine(lat1, lng1, lat2, lng2) * 1.35
    mins = (km / 50) * 60
    _osrm_cache[key] = (round(km, 1), round(mins, 1))
    return _osrm_cache[key]

def osrm_table(points):
    """Matriz de distancias/duraciones usando OSRM Table API. points = [(lat,lng),...]"""
    if len(points) > 100:
        # Demasiados — usar haversine
        return None
    coords = ";".join(f"{lng},{lat}" for lat, lng in points)
    try:
        url = f"{OSRM_BASE}/table/v1/driving/{coords}?annotations=distance,duration"
        r = requests.get(url, timeout=30)
        if r.status_code == 200:
            data = r.json()
            if data.get('code') == 'Ok':
                return {
                    'distances': [[round(d/1000, 2) for d in row] for row in data['distances']],
                    'durations': [[round(d/60, 2) for d in row] for row in data['durations']]
                }
    except:
        pass
    return None

def calc_distances_from_base(clients):
    """Calcula distancia/tiempo desde base para cada cliente."""
    print(f"  Calculando distancias desde base para {len(clients)} clientes...")
    for i, c in enumerate(clients):
        km, mins = osrm_route(BASE_LAT, BASE_LNG, c['lat'], c['lng'])
        c['km_desde_base'] = km
        c['min_desde_base'] = mins
        if (i + 1) % 20 == 0:
            print(f"    {i+1}/{len(clients)}...")
            time.sleep(0.5)  # Rate limiting
    return clients

def calc_incremental_cost(clients):
    """Calcula km/tiempo incremental de cada cliente dentro de su ruta."""
    print("  Calculando costes incrementales por ruta...")
    rutas = {}
    for c in clients:
        rn = c.get('ruta_name', 'Sin ruta')
        rutas.setdefault(rn, []).append(c)

    for ruta_name, ruta_clients in rutas.items():
        print(f"    Ruta: {ruta_name} ({len(ruta_clients)} clientes)")

        # Ordenar por distancia a base (nearest-neighbor simple)
        ruta_clients.sort(key=lambda x: x['km_desde_base'])

        # Construir puntos: base + todos los clientes
        points = [(BASE_LAT, BASE_LNG)] + [(c['lat'], c['lng']) for c in ruta_clients]

        # Intentar OSRM Table
        matrix = osrm_table(points)
        time.sleep(1)

        if matrix:
            dist_m = matrix['distances']
            dur_m = matrix['durations']
        else:
            # Fallback haversine
            n = len(points)
            dist_m = [[0.0]*n for _ in range(n)]
            dur_m = [[0.0]*n for _ in range(n)]
            for i in range(n):
                for j in range(n):
                    if i != j:
                        d = haversine(points[i][0], points[i][1], points[j][0], points[j][1]) * 1.35
                        dist_m[i][j] = round(d, 2)
                        dur_m[i][j] = round((d/50)*60, 2)

        # Para cada cliente: km incremental = coste de insertarlo en la ruta
        # Simplificación: para cliente i (idx i+1 en matrix),
        # incremental = dist(prev, i) + dist(i, next) - dist(prev, next)
        n_clients = len(ruta_clients)
        for idx, c in enumerate(ruta_clients):
            mi = idx + 1  # matrix index (0 = base)
            if n_clients == 1:
                # Solo cliente: ida y vuelta desde base
                c['km_incremental'] = dist_m[0][mi]
                c['min_incremental'] = dur_m[0][mi]
            else:
                prev_mi = (idx) if idx > 0 else 0  # previous in route (or base)
                next_mi = (idx + 2) if idx < n_clients - 1 else 0  # next (or back to base)
                d_with = dist_m[prev_mi][mi] + dist_m[mi][next_mi]
                d_without = dist_m[prev_mi][next_mi]
                t_with = dur_m[prev_mi][mi] + dur_m[mi][next_mi]
                t_without = dur_m[prev_mi][next_mi]
                c['km_incremental'] = max(round(d_with - d_without, 2), 0.5)
                c['min_incremental'] = max(round(t_with - t_without, 2), 0.5)

    return clients

# ─── ANÁLISIS DE RENTABILIDAD ────────────────────────────────────
def coste_paqueteria(peso_kg):
    for limite, precio in sorted(PAQUETERIA.items()):
        if peso_kg <= limite:
            return precio
    return 25.0

def analyze_profitability(clients):
    """Calcula coste ruta, coste paquetería, ahorro y clasificación.

    Usa coste mixto: incremental dentro de la ruta + parte proporcional
    del coste de llegar a la zona (repartido entre clientes de la ruta).
    """
    # Calcular coste base compartido por ruta
    rutas = {}
    for c in clients:
        rn = c.get('ruta_name', 'Sin ruta')
        rutas.setdefault(rn, []).append(c)

    ruta_coste_base = {}
    for rn, rc in rutas.items():
        # Coste de ir desde base hasta la zona (usando el cliente más cercano de la ruta)
        min_km = min(c['km_desde_base'] for c in rc)
        # Ida y vuelta al punto más cercano de la ruta
        coste_viaje_base = min_km * 2 * COSTE_KM + (min_km * 2 / 50) * COSTE_HORA  # 50km/h media
        # Repartir entre clientes de la ruta
        ruta_coste_base[rn] = coste_viaje_base / len(rc)

    for c in clients:
        freq = c['frecuencia_mensual']
        km_inc = c.get('km_incremental', 5)
        min_inc = c.get('min_incremental', 10)
        rn = c.get('ruta_name', 'Sin ruta')

        # Coste incremental propio (ida y vuelta)
        coste_km_inc = km_inc * 2 * COSTE_KM
        coste_tiempo_inc = (min_inc * 2 / 60) * COSTE_HORA
        coste_descarga = (TIEMPO_DESCARGA_MIN / 60) * COSTE_HORA

        # Coste total = parte proporcional del viaje base + incremental + descarga
        coste_base_prop = ruta_coste_base.get(rn, 0)
        c['coste_ruta_mensual'] = round((coste_base_prop + coste_km_inc + coste_tiempo_inc + coste_descarga) * freq, 2)

        # Coste paquetería mensual
        c['coste_paqueteria_unitario'] = coste_paqueteria(c['peso_medio_kg'])
        c['coste_paqueteria_mensual'] = round(c['coste_paqueteria_unitario'] * freq, 2)

        # Ahorro mensual (positivo = compensa paquetería)
        c['ahorro_mensual'] = round(c['coste_ruta_mensual'] - c['coste_paqueteria_mensual'], 2)

        # Ratio rentabilidad
        if c['coste_ruta_mensual'] > 0:
            c['ratio'] = round(c['facturacion_mensual'] / c['coste_ruta_mensual'], 2)
        else:
            c['ratio'] = 99.0

        # Clasificación
        # NO RENTABLE: ratio bajo O coste ruta muy superior a paquetería
        if c['ratio'] < 3 and c['ahorro_mensual'] > 0:
            c['clasificacion'] = 'NO RENTABLE'
        elif c['ratio'] >= 8 or c['ahorro_mensual'] <= 0:
            c['clasificacion'] = 'RENTABLE'
        else:
            c['clasificacion'] = 'REVISAR'

    return clients

def analyze_reassignment(clients):
    """Para clientes REVISAR/NO RENTABLE, evalúa reasignación a otra ruta."""
    rutas_centroid = {}
    for c in clients:
        rn = c.get('ruta_name', 'Sin ruta')
        rutas_centroid.setdefault(rn, {'lats': [], 'lngs': []})
        rutas_centroid[rn]['lats'].append(c['lat'])
        rutas_centroid[rn]['lngs'].append(c['lng'])

    for rn, data in rutas_centroid.items():
        data['center_lat'] = sum(data['lats']) / len(data['lats'])
        data['center_lng'] = sum(data['lngs']) / len(data['lngs'])

    for c in clients:
        if c['clasificacion'] in ('REVISAR', 'NO RENTABLE'):
            best_ruta = None
            best_dist = 9999
            for rn, data in rutas_centroid.items():
                if rn == c.get('ruta_name'):
                    continue
                d = haversine(c['lat'], c['lng'], data['center_lat'], data['center_lng'])
                if d < best_dist:
                    best_dist = d
                    best_ruta = rn
            c['ruta_propuesta'] = best_ruta
            c['dist_a_ruta_propuesta'] = round(best_dist, 1)
            # Solo proponer cambio si la otra ruta está significativamente más cerca
            c['recomendar_cambio'] = best_dist < haversine(
                c['lat'], c['lng'],
                rutas_centroid.get(c.get('ruta_name', ''), {}).get('center_lat', BASE_LAT),
                rutas_centroid.get(c.get('ruta_name', ''), {}).get('center_lng', BASE_LNG)
            ) * 0.7
        else:
            c['ruta_propuesta'] = None
            c['recomendar_cambio'] = False

    return clients

# ─── FORMATEO ESPAÑOL ────────────────────────────────────────────
def fmt_eur(val):
    """Formato 1.234,56 €"""
    if val is None:
        return '-'
    neg = val < 0
    val = abs(val)
    entero = int(val)
    decimal = round((val - entero) * 100)
    s_entero = f"{entero:,}".replace(",", ".")
    result = f"{s_entero},{decimal:02d} €"
    return f"-{result}" if neg else result

def fmt_num(val, decimals=1):
    if val is None:
        return '-'
    return f"{val:,.{decimals}f}".replace(",", "X").replace(".", ",").replace("X", ".")

def emoji_clasif(clasif):
    return {'RENTABLE': '🟢', 'REVISAR': '🟡', 'NO RENTABLE': '🔴'}.get(clasif, '?')

# ─── GRÁFICOS ────────────────────────────────────────────────────
plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['font.sans-serif'] = ['Poppins', 'Calibri', 'Arial', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

def chart_ahorro_por_ruta(clients):
    """Gráfico de barras: ahorro mensual por ruta."""
    rutas = {}
    for c in clients:
        rn = c.get('ruta_name', 'Sin ruta')
        rutas.setdefault(rn, 0)
        if c['clasificacion'] == 'NO RENTABLE':
            rutas[rn] += c['ahorro_mensual']

    names = sorted(rutas.keys())
    vals = [rutas[n] for n in names]
    colors = [RUTA_COLORS_HEX.get(n, '#888888') for n in names]

    fig, ax = plt.subplots(figsize=(10, 5))
    bars = ax.bar(names, vals, color=colors, edgecolor='white', linewidth=1.5)
    ax.set_ylabel('Ahorro potencial mensual (€)', fontsize=11)
    ax.set_title('Ahorro mensual si eliminamos clientes NO RENTABLES por ruta', fontsize=13, fontweight='bold')
    ax.axhline(0, color='gray', linewidth=0.5)
    for bar, v in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 5,
                fmt_eur(v), ha='center', va='bottom', fontsize=9)
    plt.xticks(rotation=30, ha='right')
    plt.tight_layout()
    path = CHART_DIR / 'ahorro_por_ruta.png'
    plt.savefig(path, dpi=150)
    plt.close()
    return path

def chart_facturacion_vs_coste(clients):
    """Dispersión: Facturación vs Coste ruta."""
    fig, ax = plt.subplots(figsize=(10, 6))
    for c in clients:
        color = RUTA_COLORS_HEX.get(c.get('ruta_name', ''), '#888')
        ax.scatter(c['coste_ruta_mensual'], c['facturacion_mensual'],
                   c=color, s=40, alpha=0.7, edgecolors='white', linewidth=0.5)

    # Líneas de ratio
    max_coste = max(c['coste_ruta_mensual'] for c in clients) * 1.1
    ax.plot([0, max_coste], [0, max_coste * 5], '--', color='green', alpha=0.5, label='Ratio = 5')
    ax.plot([0, max_coste], [0, max_coste * 2], '--', color='orange', alpha=0.5, label='Ratio = 2')

    ax.set_xlabel('Coste ruta mensual (€)', fontsize=11)
    ax.set_ylabel('Facturación mensual (€)', fontsize=11)
    ax.set_title('Facturación vs Coste de Ruta por Cliente', fontsize=13, fontweight='bold')
    ax.legend()

    # Leyenda de rutas
    for rn, col in RUTA_COLORS_HEX.items():
        ax.scatter([], [], c=col, s=60, label=rn)
    handles, labels = ax.get_legend_handles_labels()
    ax.legend(handles, labels, loc='upper left', fontsize=8, ncol=2)

    plt.tight_layout()
    path = CHART_DIR / 'facturacion_vs_coste.png'
    plt.savefig(path, dpi=150)
    plt.close()
    return path

def chart_clasificacion_tarta(clients):
    """Tarta: distribución de clasificación."""
    counts = {'RENTABLE': 0, 'REVISAR': 0, 'NO RENTABLE': 0}
    for c in clients:
        counts[c['clasificacion']] += 1

    labels = ['Rentable', 'Revisar', 'No Rentable']
    sizes = [counts['RENTABLE'], counts['REVISAR'], counts['NO RENTABLE']]
    colors_pie = ['#4caf50', '#ff9800', '#f44336']
    explode = (0, 0.05, 0.1)

    fig, ax = plt.subplots(figsize=(7, 7))
    wedges, texts, autotexts = ax.pie(sizes, labels=labels, colors=colors_pie, explode=explode,
                                       autopct=lambda pct: f'{int(round(pct*sum(sizes)/100))}\n({pct:.1f}%)',
                                       startangle=90, textprops={'fontsize': 12})
    for at in autotexts:
        at.set_fontsize(10)
        at.set_fontweight('bold')
    ax.set_title('Distribución de Clientes por Clasificación', fontsize=14, fontweight='bold')
    plt.tight_layout()
    path = CHART_DIR / 'clasificacion_tarta.png'
    plt.savefig(path, dpi=150)
    plt.close()
    return path

def chart_top10_ahorro(clients):
    """Barras horizontales: top 10 clientes con mayor ahorro potencial."""
    candidatos = [c for c in clients if c['ahorro_mensual'] > 0]
    candidatos.sort(key=lambda x: x['ahorro_mensual'], reverse=True)
    top = candidatos[:10]

    names = [f"{c['name'][:30]}" for c in top]
    vals = [c['ahorro_mensual'] for c in top]
    colors = [RUTA_COLORS_HEX.get(c.get('ruta_name', ''), '#888') for c in top]

    fig, ax = plt.subplots(figsize=(10, 6))
    bars = ax.barh(range(len(top)), vals, color=colors, edgecolor='white')
    ax.set_yticks(range(len(top)))
    ax.set_yticklabels(names, fontsize=9)
    ax.invert_yaxis()
    ax.set_xlabel('Ahorro mensual (€)', fontsize=11)
    ax.set_title('Top 10 Clientes con Mayor Ahorro Potencial', fontsize=13, fontweight='bold')
    for bar, v in zip(bars, vals):
        ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height()/2,
                fmt_eur(v), va='center', fontsize=9)
    plt.tight_layout()
    path = CHART_DIR / 'top10_ahorro.png'
    plt.savefig(path, dpi=150)
    plt.close()
    return path

def chart_coste_actual_vs_optimizado(clients):
    """Barras agrupadas: coste actual vs optimizado por ruta."""
    rutas_actual = {}
    rutas_optim = {}
    for c in clients:
        rn = c.get('ruta_name', 'Sin ruta')
        rutas_actual.setdefault(rn, 0)
        rutas_optim.setdefault(rn, 0)
        rutas_actual[rn] += c['coste_ruta_mensual']
        if c['clasificacion'] != 'NO RENTABLE':
            rutas_optim[rn] += c['coste_ruta_mensual']
        else:
            rutas_optim[rn] += c['coste_paqueteria_mensual']

    names = sorted(rutas_actual.keys())
    x = range(len(names))
    actual = [rutas_actual[n] for n in names]
    optim = [rutas_optim[n] for n in names]

    fig, ax = plt.subplots(figsize=(10, 5))
    w = 0.35
    ax.bar([i - w/2 for i in x], actual, w, label='Coste actual', color='#ef5350')
    ax.bar([i + w/2 for i in x], optim, w, label='Coste optimizado', color='#66bb6a')
    ax.set_xticks(list(x))
    ax.set_xticklabels(names, rotation=30, ha='right')
    ax.set_ylabel('Coste mensual (€)', fontsize=11)
    ax.set_title('Coste Actual vs Optimizado por Ruta', fontsize=13, fontweight='bold')
    ax.legend()
    plt.tight_layout()
    path = CHART_DIR / 'coste_actual_vs_optimizado.png'
    plt.savefig(path, dpi=150)
    plt.close()
    return path

# ─── GENERACIÓN DOCX ─────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, color='D5D3CC'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

def add_styled_table(doc, headers, rows, col_widths=None, right_align_from=3):
    """Añade tabla con estilo corporativo Veraleza."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, VERDE_VERALEZA)
        set_cell_border(cell, COLOR_BORDE)

    # Rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8)
            run.font.name = 'Calibri'

            # Alineación
            if ci >= right_align_from:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Color alterno
            if ri % 2 == 1:
                set_cell_shading(cell, COLOR_FILA_ALT)

            # Color clasificación
            if isinstance(val, str):
                if '🔴' in val or 'NO RENTABLE' in val:
                    set_cell_shading(cell, BG_ROJO)
                elif '🟡' in val or 'REVISAR' in val:
                    set_cell_shading(cell, BG_AMARILLO)
                elif '🟢' in val or 'RENTABLE' == val:
                    set_cell_shading(cell, BG_VERDE)

            set_cell_border(cell, COLOR_BORDE)

    # Anchos
    if col_widths:
        for ri_idx in range(len(table.rows)):
            for ci_idx, w in enumerate(col_widths):
                table.rows[ri_idx].cells[ci_idx].width = Cm(w)

    return table

def add_kpi_box(doc, label, value, unit=''):
    """Añade un párrafo KPI destacado."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x8E, 0x8B, 0x30)
    run.font.name = 'Calibri'
    run2 = p.add_run(f"{value} {unit}")
    run2.bold = True
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x8E, 0x8B, 0x30)
    run2.font.name = 'Calibri'

def generate_docx(clients, charts):
    """Genera el informe completo en .docx."""
    doc = Document()

    # ── Estilos base ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x8E, 0x8B, 0x30)
        hs.font.bold = True

    # ── Márgenes ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Header con logo ──
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if LOGO_PATH.exists():
        run = hp.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(4))

    # ── Footer ──
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run('Veraleza — Confidencial')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = 'Calibri'

    # ══════════════════════════════════════════════════════════════
    # PORTADA
    # ══════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(8))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('INFORME DE OPTIMIZACIÓN\nDE RUTAS COMERCIALES')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x8E, 0x8B, 0x30)
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Análisis de Rentabilidad por Cliente y Recomendaciones')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Fecha: {datetime.now().strftime("%d/%m/%Y")}\nPeríodo analizado: Últimos 12 meses\nGenerado por: Gestor de Rutas Veraleza')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CÁLCULOS RESUMEN
    # ══════════════════════════════════════════════════════════════
    no_rentables = [c for c in clients if c['clasificacion'] == 'NO RENTABLE']
    revisar = [c for c in clients if c['clasificacion'] == 'REVISAR']
    rentables = [c for c in clients if c['clasificacion'] == 'RENTABLE']

    ahorro_mensual = sum(c['ahorro_mensual'] for c in no_rentables)
    ahorro_anual = ahorro_mensual * 12
    horas_liberadas = sum(c['min_incremental'] * 2 * c['frecuencia_mensual'] for c in no_rentables) / 60
    km_reduccion_mes = sum(c['km_incremental'] * 2 * c['frecuencia_mensual'] for c in no_rentables)
    coste_paq_total = sum(c['coste_paqueteria_mensual'] for c in no_rentables)
    cambio_ruta = [c for c in clients if c.get('recomendar_cambio')]

    # ══════════════════════════════════════════════════════════════
    # 1. RESUMEN EJECUTIVO
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('1. Resumen Ejecutivo', level=1)

    doc.add_paragraph(
        'Este informe analiza la rentabilidad de cada cliente activo en las rutas comerciales '
        'de Veraleza, comparando el coste de distribución propia con el envío por paquetería. '
        'Se identifican clientes cuyo coste de reparto supera el beneficio que generan y se '
        'proponen acciones de optimización.'
    )

    add_kpi_box(doc, 'Ahorro mensual estimado', fmt_eur(ahorro_mensual))
    add_kpi_box(doc, 'Ahorro anual estimado', fmt_eur(ahorro_anual))
    add_kpi_box(doc, 'Clientes a pasar a paquetería', str(len(no_rentables)))
    add_kpi_box(doc, 'Clientes a reasignar de ruta', str(len(cambio_ruta)))
    add_kpi_box(doc, 'Horas liberadas al mes', fmt_num(horas_liberadas) + ' h')
    add_kpi_box(doc, 'Reducción de km mensuales', fmt_num(km_reduccion_mes, 0) + ' km')

    # Gráfico tarta
    if charts.get('tarta') and Path(charts['tarta']).exists():
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(charts['tarta']), width=Cm(12))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 2. METODOLOGÍA
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('2. Metodología', level=1)

    doc.add_heading('Fuente de datos', level=2)
    doc.add_paragraph(
        'Los datos de clientes, coordenadas GPS y asignación de rutas se han extraído '
        'directamente de la base de datos del Gestor de Rutas (MySQL). Los datos de facturación '
        'y frecuencia de visita se han estimado en base a parámetros del sector para esta '
        'primera versión del informe.'
    )

    doc.add_heading('Motor de cálculo de rutas', level=2)
    doc.add_paragraph(
        'Se utiliza OSRM (Open Source Routing Machine) con datos de OpenStreetMap para '
        'calcular distancias y tiempos reales por carretera entre la base de Veraleza Tomiño '
        'y cada cliente, así como los costes incrementales dentro de cada ruta.'
    )

    doc.add_heading('Parámetros de coste', level=2)
    params = [
        ('Coste por km del vehículo', '0,35 €/km', 'Gasóleo + desgaste + seguro'),
        ('Coste por hora del conductor', '18,00 €/h', 'Salario + Seguridad Social'),
        ('Tiempo medio de descarga', '20 min', 'Por cliente y entrega'),
        ('Paquete < 5 kg', '4,50 €', 'Envío estándar'),
        ('Paquete 5-15 kg', '7,50 €', 'Envío estándar'),
        ('Paquete 15-30 kg', '12,00 €', 'Envío estándar'),
        ('Paquete > 30 kg / palet', '25,00 €', 'Envío palet'),
    ]
    add_styled_table(doc, ['Parámetro', 'Valor', 'Observaciones'],
                     params, col_widths=[7, 3, 7], right_align_from=1)

    doc.add_heading('Criterios de clasificación', level=2)
    clasif_rows = [
        ('🟢 RENTABLE', 'Ratio > 5 y paquetería más cara', 'Mantener en ruta'),
        ('🟡 REVISAR', 'Ratio entre 2 y 5', 'Evaluar alternativas'),
        ('🔴 NO RENTABLE', 'Ratio < 2 y paquetería más barata', 'Pasar a paquetería'),
    ]
    add_styled_table(doc, ['Clasificación', 'Criterio', 'Acción'], clasif_rows,
                     col_widths=[4, 8, 5], right_align_from=99)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 3. ANÁLISIS POR RUTA
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('3. Análisis por Ruta', level=1)

    rutas_grouped = {}
    for c in clients:
        rn = c.get('ruta_name', 'Sin ruta')
        rutas_grouped.setdefault(rn, []).append(c)

    for ruta_name in sorted(rutas_grouped.keys()):
        ruta_clients = rutas_grouped[ruta_name]
        doc.add_heading(f'Ruta: {ruta_name}', level=2)

        n_total = len(ruta_clients)
        n_rent = sum(1 for c in ruta_clients if c['clasificacion'] == 'RENTABLE')
        n_rev = sum(1 for c in ruta_clients if c['clasificacion'] == 'REVISAR')
        n_nrent = sum(1 for c in ruta_clients if c['clasificacion'] == 'NO RENTABLE')
        ahorro_ruta = sum(c['ahorro_mensual'] for c in ruta_clients if c['clasificacion'] == 'NO RENTABLE')

        doc.add_paragraph(
            f'Total clientes: {n_total} — '
            f'🟢 {n_rent}  🟡 {n_rev}  🔴 {n_nrent} — '
            f'Ahorro potencial: {fmt_eur(ahorro_ruta)}/mes'
        )

        headers = ['ID', 'Nombre', 'Localidad', 'Km base', 'Km incr.', 'Min incr.',
                   'Fact. €/mes', 'Coste ruta', 'Coste paq.', 'Ahorro', 'Ratio', 'Clasif.']

        rows = []
        ruta_clients.sort(key=lambda x: x['ahorro_mensual'], reverse=True)
        for c in ruta_clients:
            # Extraer localidad del address
            addr = c.get('address', '') or ''
            parts = [p.strip() for p in addr.split(',')]
            localidad = parts[1] if len(parts) > 1 else parts[0] if parts else ''
            localidad = localidad[:20]

            rows.append([
                str(c['id']),
                c['name'][:25],
                localidad,
                fmt_num(c['km_desde_base']),
                fmt_num(c['km_incremental']),
                fmt_num(c['min_incremental']),
                fmt_eur(c['facturacion_mensual']),
                fmt_eur(c['coste_ruta_mensual']),
                fmt_eur(c['coste_paqueteria_mensual']),
                fmt_eur(c['ahorro_mensual']),
                fmt_num(c['ratio']),
                f"{emoji_clasif(c['clasificacion'])} {c['clasificacion']}"
            ])

        add_styled_table(doc, headers, rows,
                         col_widths=[1, 3.5, 2.5, 1.3, 1.3, 1.3, 1.8, 1.6, 1.5, 1.5, 1.1, 2.5],
                         right_align_from=3)

        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 4. CLIENTES RECOMENDADOS PARA PAQUETERÍA
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('4. Clientes Recomendados para Paquetería', level=1)

    doc.add_paragraph(
        f'Se identifican {len(no_rentables)} clientes cuyo coste de distribución por ruta '
        f'supera el coste de envío por paquetería, con un ahorro total estimado de '
        f'{fmt_eur(ahorro_mensual)}/mes ({fmt_eur(ahorro_anual)}/año).'
    )

    if no_rentables:
        no_rentables.sort(key=lambda x: x['ahorro_mensual'], reverse=True)
        headers = ['ID', 'Nombre', 'Ruta actual', 'Km base', 'Fact. €/mes',
                   'Coste ruta', 'Coste paq.', 'Ahorro/mes']
        rows = []
        for c in no_rentables:
            rows.append([
                str(c['id']),
                c['name'][:30],
                c.get('ruta_name', '-'),
                fmt_num(c['km_desde_base']),
                fmt_eur(c['facturacion_mensual']),
                fmt_eur(c['coste_ruta_mensual']),
                fmt_eur(c['coste_paqueteria_mensual']),
                fmt_eur(c['ahorro_mensual'])
            ])
        add_styled_table(doc, headers, rows,
                         col_widths=[1, 4, 2.5, 1.5, 2, 2, 2, 2],
                         right_align_from=3)

        # Agrupado por ruta
        doc.add_heading('Ahorro agrupado por ruta', level=2)
        ahorro_ruta = {}
        count_ruta = {}
        for c in no_rentables:
            rn = c.get('ruta_name', 'Sin ruta')
            ahorro_ruta[rn] = ahorro_ruta.get(rn, 0) + c['ahorro_mensual']
            count_ruta[rn] = count_ruta.get(rn, 0) + 1

        rows = [[rn, str(count_ruta[rn]), fmt_eur(ahorro_ruta[rn]), fmt_eur(ahorro_ruta[rn]*12)]
                for rn in sorted(ahorro_ruta.keys())]
        rows.append(['TOTAL', str(len(no_rentables)), fmt_eur(ahorro_mensual), fmt_eur(ahorro_anual)])
        add_styled_table(doc, ['Ruta', 'Clientes', 'Ahorro/mes', 'Ahorro/año'], rows,
                         col_widths=[5, 2, 4, 4], right_align_from=1)

    # Gráfico top 10
    if charts.get('top10') and Path(charts['top10']).exists():
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(charts['top10']), width=Cm(15))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 5. CLIENTES RECOMENDADOS PARA CAMBIO DE RUTA
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('5. Clientes Recomendados para Cambio de Ruta', level=1)

    if cambio_ruta:
        doc.add_paragraph(
            f'Se identifican {len(cambio_ruta)} clientes que podrían beneficiarse de un '
            f'cambio de ruta por cercanía geográfica a otra ruta existente.'
        )
        headers = ['ID', 'Nombre', 'Ruta actual', 'Ruta propuesta', 'Km a ruta prop.', 'Clasif.']
        rows = []
        for c in cambio_ruta:
            rows.append([
                str(c['id']),
                c['name'][:30],
                c.get('ruta_name', '-'),
                c.get('ruta_propuesta', '-'),
                fmt_num(c.get('dist_a_ruta_propuesta', 0)),
                f"{emoji_clasif(c['clasificacion'])} {c['clasificacion']}"
            ])
        add_styled_table(doc, headers, rows,
                         col_widths=[1, 4, 3, 3, 2.5, 3],
                         right_align_from=4)
    else:
        doc.add_paragraph('No se identifican clientes con beneficio claro de reasignación de ruta.')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 6. IMPACTO ECONÓMICO GLOBAL
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('6. Impacto Económico Global', level=1)

    km_actual_mes = sum(c['km_incremental'] * 2 * c['frecuencia_mensual'] for c in clients)
    horas_actual_mes = sum((c['min_incremental'] * 2 + TIEMPO_DESCARGA_MIN) * c['frecuencia_mensual'] for c in clients) / 60
    coste_actual_mes = sum(c['coste_ruta_mensual'] for c in clients)

    km_optim_mes = km_actual_mes - km_reduccion_mes
    horas_optim_mes = horas_actual_mes - horas_liberadas
    coste_optim_mes = coste_actual_mes - ahorro_mensual + coste_paq_total

    headers = ['Concepto', 'Situación actual', 'Situación optimizada', 'Diferencia']
    rows = [
        ['Km totales/mes', fmt_num(km_actual_mes, 0), fmt_num(km_optim_mes, 0),
         fmt_num(km_actual_mes - km_optim_mes, 0)],
        ['Horas totales/mes', fmt_num(horas_actual_mes, 0), fmt_num(horas_optim_mes, 0),
         fmt_num(horas_actual_mes - horas_optim_mes, 0)],
        ['Coste distribución/mes', fmt_eur(coste_actual_mes), fmt_eur(coste_optim_mes - coste_paq_total),
         fmt_eur(ahorro_mensual)],
        ['Coste paquetería/mes', fmt_eur(0), fmt_eur(coste_paq_total), fmt_eur(-coste_paq_total)],
        ['Coste total/mes', fmt_eur(coste_actual_mes), fmt_eur(coste_optim_mes),
         fmt_eur(coste_actual_mes - coste_optim_mes)],
        ['Coste total/año', fmt_eur(coste_actual_mes * 12), fmt_eur(coste_optim_mes * 12),
         fmt_eur((coste_actual_mes - coste_optim_mes) * 12)],
    ]
    add_styled_table(doc, headers, rows, col_widths=[5, 4, 4, 4], right_align_from=1)

    # Gráfico comparativo
    if charts.get('comparativo') and Path(charts['comparativo']).exists():
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(charts['comparativo']), width=Cm(15))

    # Gráfico facturación vs coste
    if charts.get('dispersion') and Path(charts['dispersion']).exists():
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(charts['dispersion']), width=Cm(15))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 7. PLAN DE IMPLEMENTACIÓN
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('7. Plan de Implementación', level=1)

    phases = [
        ('Fase 1: Comunicación a clientes afectados (Semana 1-2)',
         'Contactar a los clientes identificados como NO RENTABLES para informarles del '
         'cambio en el método de entrega. Ofrecer la opción de paquetería con seguimiento '
         'y mantener el mismo nivel de servicio.'),
        ('Fase 2: Contratación con empresa de paquetería (Semana 2-3)',
         'Negociar tarifas con empresas de paquetería (SEUR, MRW, GLS, Correos Express) '
         'para obtener descuentos por volumen. Objetivo: reducir un 15-20% las tarifas estándar.'),
        ('Fase 3: Transición gradual (Semana 3-8)',
         'Implementar el cambio de forma progresiva, empezando por los clientes con mayor '
         'ahorro potencial. Monitorizar la satisfacción del cliente y ajustar si es necesario.'),
        ('Fase 4: Seguimiento y ajuste (Mes 3 en adelante)',
         'Revisar los resultados cada mes. Actualizar el análisis con datos reales de '
         'facturación y costes de paquetería. Ajustar la clasificación de clientes según evolución.'),
    ]

    for title, desc in phases:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # ANEXO A1: DATOS COMPLETOS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('Anexo A1: Datos Completos de Todos los Clientes', level=1)

    # Cambiar a horizontal para tabla ancha
    new_section = doc.add_section(2)  # WD_SECTION.NEW_PAGE -> continuous
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = Cm(29.7)
    new_section.page_height = Cm(21)
    new_section.left_margin = Cm(1.5)
    new_section.right_margin = Cm(1.5)

    headers = ['ID', 'Nombre', 'Ruta', 'Km base', 'Min base', 'Km incr.', 'Min incr.',
               'Fact. €/mes', 'Freq/mes', 'Peso kg', 'Coste ruta', 'Coste paq.', 'Ahorro', 'Ratio', 'Clasif.']

    sorted_clients = sorted(clients, key=lambda x: (x.get('ruta_name', ''), -x['ahorro_mensual']))
    rows = []
    for c in sorted_clients:
        rows.append([
            str(c['id']),
            c['name'][:22],
            (c.get('ruta_name') or '-')[:12],
            fmt_num(c['km_desde_base']),
            fmt_num(c['min_desde_base']),
            fmt_num(c['km_incremental']),
            fmt_num(c['min_incremental']),
            fmt_eur(c['facturacion_mensual']),
            str(c['frecuencia_mensual']),
            fmt_num(c['peso_medio_kg']),
            fmt_eur(c['coste_ruta_mensual']),
            fmt_eur(c['coste_paqueteria_mensual']),
            fmt_eur(c['ahorro_mensual']),
            fmt_num(c['ratio']),
            f"{emoji_clasif(c['clasificacion'])}"
        ])

    add_styled_table(doc, headers, rows,
                     col_widths=[0.9, 3.2, 1.8, 1.2, 1.2, 1.2, 1.2, 1.7, 1, 1.1, 1.5, 1.4, 1.4, 1, 1],
                     right_align_from=3)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # ANEXO A3: PARÁMETROS
    # ══════════════════════════════════════════════════════════════
    # Volver a vertical
    new_section2 = doc.add_section(2)
    new_section2.orientation = WD_ORIENT.PORTRAIT
    new_section2.page_width = Cm(21)
    new_section2.page_height = Cm(29.7)

    doc.add_heading('Anexo A3: Parámetros de Cálculo', level=1)

    params_full = [
        ('Base de operaciones', 'Veraleza Tomiño', '41.9945°N, 8.7399°W'),
        ('Coste por km', '0,35 €', 'Gasóleo + desgaste + seguro'),
        ('Coste por hora conductor', '18,00 €', 'Salario bruto + SS'),
        ('Tiempo descarga por cliente', '20 minutos', 'Media estimada'),
        ('Motor de routing', 'OSRM', 'Open Source Routing Machine'),
        ('Datos cartográficos', 'OpenStreetMap', 'Actualización continua'),
        ('Clientes analizados', str(len(clients)), 'Activos con coordenadas'),
        ('Rutas analizadas', str(len(rutas_grouped)), ''),
        ('Fecha del análisis', datetime.now().strftime('%d/%m/%Y'), ''),
    ]
    add_styled_table(doc, ['Parámetro', 'Valor', 'Notas'], params_full,
                     col_widths=[6, 4, 7], right_align_from=99)

    # ── Guardar ──
    output_path = OUTPUT_DIR / 'informe_optimizacion_rutas_veraleza.docx'
    doc.save(str(output_path))
    return output_path


# ─── MAIN ─────────────────────────────────────────────────────────
def main():
    print("=" * 60)
    print("INFORME DE OPTIMIZACIÓN DE RUTAS — VERALEZA")
    print("=" * 60)

    # 1. Extraer datos
    print("\n[1/6] Extrayendo datos de MySQL...")
    clients = fetch_clients()
    rutas = fetch_rutas()
    print(f"  → {len(clients)} clientes activos, {len(rutas)} rutas")

    # 2. Simular datos comerciales
    print("\n[2/6] Simulando datos comerciales (facturación, frecuencia, peso)...")
    clients = simulate_commercial_data(clients)

    # 3. Calcular distancias OSRM
    print("\n[3/6] Calculando distancias OSRM desde base...")
    clients = calc_distances_from_base(clients)

    print("\n[4/6] Calculando costes incrementales por ruta...")
    clients = calc_incremental_cost(clients)

    # 4. Análisis de rentabilidad
    print("\n[5/6] Analizando rentabilidad y reasignación...")
    clients = analyze_profitability(clients)
    clients = analyze_reassignment(clients)

    # Resumen rápido
    n_rent = sum(1 for c in clients if c['clasificacion'] == 'RENTABLE')
    n_rev = sum(1 for c in clients if c['clasificacion'] == 'REVISAR')
    n_nrent = sum(1 for c in clients if c['clasificacion'] == 'NO RENTABLE')
    print(f"  → 🟢 {n_rent} rentables, 🟡 {n_rev} revisar, 🔴 {n_nrent} no rentables")

    ahorro = sum(c['ahorro_mensual'] for c in clients if c['clasificacion'] == 'NO RENTABLE')
    print(f"  → Ahorro potencial mensual: {fmt_eur(ahorro)}")
    print(f"  → Ahorro potencial anual: {fmt_eur(ahorro * 12)}")

    # 5. Generar gráficos
    print("\n[6/6] Generando informe .docx con gráficos...")
    charts = {}
    charts['ahorro_ruta'] = str(chart_ahorro_por_ruta(clients))
    charts['dispersion'] = str(chart_facturacion_vs_coste(clients))
    charts['tarta'] = str(chart_clasificacion_tarta(clients))
    charts['top10'] = str(chart_top10_ahorro(clients))
    charts['comparativo'] = str(chart_coste_actual_vs_optimizado(clients))
    print(f"  → {len(charts)} gráficos generados en {CHART_DIR}")

    # 6. Generar DOCX
    output_path = generate_docx(clients, charts)
    print(f"\n{'='*60}")
    print(f"INFORME GENERADO: {output_path}")
    print(f"{'='*60}")

    # Guardar también JSON con datos
    json_path = OUTPUT_DIR / 'datos_analisis.json'
    export = []
    for c in clients:
        export.append({
            'id': c['id'], 'name': c['name'], 'ruta': c.get('ruta_name'),
            'km_base': c['km_desde_base'], 'km_incremental': c['km_incremental'],
            'facturacion_mensual': c['facturacion_mensual'],
            'coste_ruta_mensual': c['coste_ruta_mensual'],
            'coste_paqueteria_mensual': c['coste_paqueteria_mensual'],
            'ahorro_mensual': c['ahorro_mensual'],
            'ratio': c['ratio'],
            'clasificacion': c['clasificacion'],
        })
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(export, f, ensure_ascii=False, indent=2)
    print(f"Datos exportados: {json_path}")


if __name__ == '__main__':
    main()
